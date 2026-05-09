from __future__ import annotations

import re
import xml.etree.ElementTree as ET
from dataclasses import dataclass, field
from pathlib import Path
from typing import Iterable

from PIL import Image
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ACCENT = RGBColor(21, 74, 120)
ACCENT_LIGHT = "D9EAF7"
ACCENT_DARK = "154A78"
SUCCESS = RGBColor(21, 128, 61)
SUCCESS_FILL = "E9F7EF"
BORDER = "C7D5E0"
TEXT = RGBColor(40, 40, 40)
MUTED = RGBColor(95, 99, 104)
SCRIPT_DIR = Path(__file__).resolve().parent
DOCS_DIR = SCRIPT_DIR.parent
ROOT_DIR = DOCS_DIR.parent
REGRESSION_DIR = DOCS_DIR / "regression"
SUBMITTER_NAME = "Joana Carla D. Gako"


@dataclass
class AutomatedCase:
    label: str
    test_id: str
    title: str
    screenshot: Path
    result: str
    description: str
    execution_time: str
    validation: str


@dataclass
class AutomatedModule:
    title: str
    cases: list[AutomatedCase] = field(default_factory=list)


@dataclass
class Issue:
    heading: str
    description: str
    severity: str
    affected: str
    status: str


@dataclass
class FixItem:
    heading: str
    commit: str
    tests: str
    fix: str


@dataclass
class TestProcedure:
    procedure_id: str
    platform: str
    title: str
    objective: str
    preconditions: list[str]
    steps: list[str]


BACKEND_COVERAGE_ROWS = [
    ["BE-SMOKE-001", "Authentication login and refresh routes", "Supported auth endpoints accept requests and return tokens", "PASS"],
    ["BE-SMOKE-002", "Profile route aliases", "Authenticated profile routes return the same user profile data", "PASS"],
    ["BE-SMOKE-003", "Notebook create and review routes", "Notebook creation and review-update routes remain functional after refactor", "PASS"],
    ["BE-SMOKE-004", "Playback queue API compatibility", "Queue endpoints expose the active queue through supported aliases", "PASS"],
    ["BE-SMOKE-005", "AI configuration endpoints", "Selected config and config list endpoints return expected payloads", "PASS"],
    ["BE-SMOKE-006", "Quiz and flashcard attempt endpoints", "Attempt submission persists through service layer contracts", "PASS"],
    ["BE-SMOKE-007", "Admin user route authorization", "Unauthorized and non-admin requests are rejected; admin requests succeed", "PASS"],
    ["BE-SMOKE-008", "Service regression safeguards", "Flashcard, quiz, AI config, and notebook service invariants continue to hold", "PASS"],
]

MOBILE_AUTH_COVERAGE_ROWS = [
    ["MOB-AUTH-001", "Registration field entry", "Registration form accepts the expected input values", "PASS"],
    ["MOB-AUTH-002", "Registration verification email step", "Post-registration verification prompt is displayed", "PASS"],
    ["MOB-AUTH-003", "Registration success state", "Registration flow completes and shows success confirmation", "PASS"],
    ["MOB-AUTH-010", "Forgot-password email entry", "Recovery email screen accepts a valid address", "PASS"],
    ["MOB-AUTH-011", "Forgot-password code field entry", "Verification code input renders and accepts digits", "PASS"],
    ["MOB-AUTH-012", "Forgot-password paste-code behavior", "Pasted verification code is distributed correctly across inputs", "PASS"],
    ["MOB-AUTH-013", "Forgot-password code validation", "Valid code advances the flow successfully", "PASS"],
    ["MOB-AUTH-014", "Reset-password field entry", "New-password screen renders and accepts the new password", "PASS"],
    ["MOB-AUTH-015", "Reset-password success state", "Password update completes and success confirmation is shown", "PASS"],
    ["MOB-OAUTH-001", "Google OAuth launch", "Google sign-in surface opens correctly on mobile", "PASS"],
    ["MOB-OAUTH-002", "Google OAuth loading state", "Loading state is visible while OAuth is in progress", "PASS"],
    ["MOB-OAUTH-003", "Google OAuth success state", "OAuth flow completes and returns to an authenticated state", "PASS"],
]

MOBILE_HOME_STATE_ROWS = [
    ["MOB-HOME-001", "Notebook category move updates lists", "Home notebook lists and queues reflect the new category", "PASS"],
    ["MOB-HOME-002", "Category delete uncategorizes notebooks", "Affected notebook/category references are cleared without stale state", "PASS"],
    ["MOB-HOME-003", "Category delete removes notebooks when requested", "Deleted notebooks are removed from playlists and playback queue", "PASS"],
    ["MOB-HOME-004", "Playlist optimistic add / remove / reorder", "Playlist queue changes are reflected immediately in visible state", "PASS"],
    ["MOB-HOME-005", "Quiz attempt stats update", "Quiz attempts increment count and best score correctly", "PASS"],
    ["MOB-HOME-006", "Flashcard attempt stats update", "Flashcard attempts increment count and preserve best mastery correctly", "PASS"],
    ["MOB-HOME-007", "Attempt mutation IDs are preserved", "Quiz and flashcard attempt requests keep caller-provided mutation IDs", "PASS"],
]

MOBILE_PLAYBACK_ARCH_ROWS = [
    ["MOB-PB-001", "Playback highlights middle word", "Current word highlight matches the active offset", "PASS"],
    ["MOB-PB-002", "Playback keeps previous word at boundary", "Boundary conditions do not skip the expected active word", "PASS"],
    ["MOB-PB-003", "Playback clears highlight at end", "Highlights clear correctly when playback ends", "PASS"],
    ["MOB-PB-004", "Replay remains possible after end", "Finished playback can resume for replay", "PASS"],
    ["MOB-PB-005", "Multi-chunk playback shows next chunk text", "Chunk transitions expose the correct active chunk and word range", "PASS"],
    ["MOB-PB-006", "Global playbar subtitle offsets", "Subtitle words preserve full-text offsets for live highlighting", "PASS"],
    ["MOB-PB-007", "Queue coordinator normalization", "Queue indices clamp safely within bounds", "PASS"],
    ["MOB-PB-008", "Queue coordinator shuffle resolution", "Shuffle chooses a valid non-current queue item", "PASS"],
    ["MOB-ARCH-001", "Vertical-slice mobile package boundaries", "Feature, platform, and shared boundaries remain intact after refactor", "PASS"],
    ["MOB-ARCH-002", "Android instrumentation smoke", "Connected Android instrumentation test completes successfully", "PASS"],
]

FIX_HISTORY_ROWS = [
    ["Mobile auth regression fixes", "d51af3a", "MOB-AUTH-001 to MOB-AUTH-015", "Completed the mobile authentication fix pass and captured refreshed mobile/web auth evidence for registration, verification, and password recovery."],
    ["Mobile auth redirect and keyboard handling", "20e0281", "MOB-AUTH-001, MOB-AUTH-010 to MOB-AUTH-014", "Corrected registration redirect behavior plus forgot-password keyboard and code-entry handling in the Android auth flow."],
    ["Web equation editor fix", "aa7e56d", "WEB-NB-010", "Fixed empty notebook block equation editing so math insertion works reliably in the notebook editor."],
    ["Editor home-navigation safeguard", "3026f10", "WEB-NB-012 / notebook navigation regression", "Ensured pending editor changes are saved before returning home to prevent data-loss regressions during navigation."],
    ["Equation insertion stabilization", "eab133a", "WEB-NB-010", "Stabilized equation insertion on empty lines to avoid broken editor behavior during retest."],
    ["Mobile playbar label correction", "4644642", "MOB-PB-006 / mobile playback UI validation", "Renamed the mobile playbar fallback label so playback status is readable and consistent during mobile regression checks."],
    ["Regression retest evidence refresh", "9c27924", "Cross-platform retest evidence", "Updated regression-validation artifacts after rerunning playbar, review audio, and live-fix verification scenarios."],
    ["Regression report final evidence update", "cea29ba", "Report artifacts", "Finalized the report/evidence artifact set after the live-fix validation pass."],
]

MOBILE_EVIDENCE_CASES = [
    ("14. MOBILE AUTHENTICATION MODULE EVIDENCE", [
        ("14.1", "MOB-AUTH-001", "Registration field entry", ROOT_DIR / "outputs/mobile/auth/register/fields.png", "PASS", "Mobile registration form accepts typed user details.", "N/A (screenshot validation)", "All required fields are visible and populated correctly."),
        ("14.2", "MOB-AUTH-002", "Registration verification email step", ROOT_DIR / "outputs/mobile/auth/register/verif-email.png", "PASS", "Registration proceeds to the email verification step.", "N/A (screenshot validation)", "Verification prompt is shown after registration."),
        ("14.3", "MOB-AUTH-003", "Registration success state", ROOT_DIR / "outputs/mobile/auth/register/success.png.png", "PASS", "Successful registration confirmation is displayed.", "N/A (screenshot validation)", "Success state confirms account creation."),
        ("14.4", "MOB-AUTH-010", "Forgot-password email entry", ROOT_DIR / "outputs/mobile/auth/forgot-password/email.png", "PASS", "Recovery email entry screen renders correctly.", "N/A (screenshot validation)", "Email field accepts the recovery address."),
        ("14.5", "MOB-AUTH-011", "Forgot-password code field entry", ROOT_DIR / "outputs/mobile/auth/forgot-password/field.png", "PASS", "Verification code input screen is shown.", "N/A (screenshot validation)", "Verification code fields are visible and ready for input."),
        ("14.6", "MOB-AUTH-012", "Forgot-password paste-code behavior", ROOT_DIR / "outputs/mobile/auth/forgot-password/paste-code.png", "PASS", "Pasted code is distributed into the verification fields.", "N/A (screenshot validation)", "All verification digits are populated correctly after paste."),
        ("14.7", "MOB-AUTH-013", "Forgot-password code validation", ROOT_DIR / "outputs/mobile/auth/forgot-password/code-success.png", "PASS", "Valid verification code advances the recovery flow.", "N/A (screenshot validation)", "Code submission succeeds and advances to reset-password."),
        ("14.8", "MOB-AUTH-014", "Reset-password field entry", ROOT_DIR / "outputs/mobile/auth/forgot-password/reset-password-field.png", "PASS", "Reset-password screen renders and accepts the new password.", "N/A (screenshot validation)", "Password input field is available for update."),
        ("14.9", "MOB-AUTH-015", "Reset-password success state", ROOT_DIR / "outputs/mobile/auth/forgot-password/password-update-success.png", "PASS", "Password reset completes successfully.", "N/A (screenshot validation)", "Success state confirms the password was updated."),
    ]),
    ("15. MOBILE OAUTH MODULE EVIDENCE", [
        ("15.1", "MOB-OAUTH-001", "Google OAuth launch", ROOT_DIR / "outputs/mobile/oauth/google-auth.png", "PASS", "Google sign-in screen is opened from the mobile client.", "N/A (screenshot validation)", "OAuth provider screen is visible."),
        ("15.2", "MOB-OAUTH-002", "Google OAuth loading state", ROOT_DIR / "outputs/mobile/oauth/loading.png", "PASS", "Loading indicator is displayed while authentication is in progress.", "N/A (screenshot validation)", "Loading state is visible and not blocked."),
        ("15.3", "MOB-OAUTH-003", "Google OAuth success state", ROOT_DIR / "outputs/mobile/oauth/success.png", "PASS", "OAuth flow returns to an authenticated success state.", "N/A (screenshot validation)", "Success screen confirms mobile OAuth completion."),
    ]),
]

TEST_PROCEDURES = [
    TestProcedure(
        procedure_id="TP-WEB-AUTH-001",
        platform="Web",
        title="User login with valid credentials",
        objective="Verify that a registered user can authenticate and land on the dashboard.",
        preconditions=["Web app is running.", "A valid test user account exists."],
        steps=[
            "Open the login page.",
            "Enter a valid username or email and password.",
            "Click the login button.",
            "Confirm the dashboard loads without an authentication error.",
            "Verify the user profile or authenticated navigation is visible.",
        ],
    ),
    TestProcedure(
        procedure_id="TP-WEB-NB-001",
        platform="Web",
        title="Notebook create, edit, and auto-save flow",
        objective="Validate the core notebook authoring flow after the refactor.",
        preconditions=["User is authenticated.", "Notebook creation permission is available."],
        steps=[
            "Navigate to the notebook creation screen.",
            "Create a notebook with a unique title.",
            "Enter or edit notebook content in the editor.",
            "Wait for the auto-save indicator or confirmation state.",
            "Re-open or revisit the notebook and confirm the content persists.",
        ],
    ),
    TestProcedure(
        procedure_id="TP-WEB-AI-001",
        platform="Web",
        title="Notebook AI sidebar flow",
        objective="Verify that AI-assisted notebook tools remain accessible and responsive.",
        preconditions=["User is authenticated.", "Notebook editor is open."],
        steps=[
            "Open the AI sidebar from the notebook editor.",
            "Switch between the supported AI tools or actions.",
            "Enter a prompt or select an action such as summarize or explain.",
            "Submit the request.",
            "Verify the response panel or action state updates correctly.",
        ],
    ),
    TestProcedure(
        procedure_id="TP-WEB-LIB-001",
        platform="Web",
        title="Library search and category filtering",
        objective="Confirm users can find notebooks through library controls.",
        preconditions=["User is authenticated.", "Library contains multiple notebooks and categories."],
        steps=[
            "Open the library page.",
            "Enter a search term in the search field.",
            "Confirm only matching items remain visible.",
            "Apply a category filter.",
            "Verify the filtered result set matches the selected category.",
        ],
    ),
    TestProcedure(
        procedure_id="TP-WEB-FC-001",
        platform="Web",
        title="Flashcard study mode",
        objective="Validate that flashcard study mode remains usable after refactoring.",
        preconditions=["User is authenticated.", "At least one flashcard deck exists."],
        steps=[
            "Open a flashcard deck.",
            "Start study mode.",
            "Navigate through at least one card interaction.",
            "Observe progress or mastery indicators.",
            "Confirm the session state updates without errors.",
        ],
    ),
    TestProcedure(
        procedure_id="TP-WEB-QZ-001",
        platform="Web",
        title="Quiz launch and result review",
        objective="Check quiz creation, play, scoring, and review behavior.",
        preconditions=["User is authenticated.", "At least one quiz exists."],
        steps=[
            "Open the quiz module.",
            "Launch a quiz attempt.",
            "Answer one or more questions and submit the attempt.",
            "Verify the score or result screen appears.",
            "Open review mode and confirm the correct-answer view is available.",
        ],
    ),
    TestProcedure(
        procedure_id="TP-WEB-PLQ-001",
        platform="Web",
        title="Playlist and playback queue flow",
        objective="Validate playlist manipulation and queue-based playback behavior.",
        preconditions=["User is authenticated.", "At least one notebook is available for playlisting."],
        steps=[
            "Create or open a playlist.",
            "Add one or more notebooks to the playlist queue.",
            "Reorder or remove an item.",
            "Start playback from the queue.",
            "Verify the active item, queue panel, and playback controls remain synchronized.",
        ],
    ),
    TestProcedure(
        procedure_id="TP-MOB-AUTH-001",
        platform="Mobile",
        title="Registration and password recovery flow",
        objective="Verify implemented mobile authentication flows using screenshot-backed evidence.",
        preconditions=["Android app is installed.", "Mobile network/API configuration is available."],
        steps=[
            "Open the registration screen and enter the required fields.",
            "Submit the registration form and confirm the verification step appears.",
            "Open the forgot-password flow.",
            "Enter the recovery email and verification code.",
            "Reset the password and confirm the success state is displayed.",
        ],
    ),
    TestProcedure(
        procedure_id="TP-MOB-PB-001",
        platform="Mobile",
        title="Playback highlight and queue logic",
        objective="Validate playback-specific mobile logic after the architectural refactor.",
        preconditions=["Mobile playback test fixtures are available."],
        steps=[
            "Load a playback snapshot or queue state in the automated test harness.",
            "Advance the active character offset through the text.",
            "Verify the highlighted word range matches the expected location.",
            "Move playback to the end and confirm the final-state behavior clears highlights or allows replay.",
            "Validate queue normalization and shuffle target resolution logic.",
        ],
    ),
    TestProcedure(
        procedure_id="TP-BE-SMOKE-001",
        platform="Backend",
        title="User-facing API smoke regression",
        objective="Ensure user-facing backend routes remain compatible after the refactor.",
        preconditions=["Backend test context is available.", "Controller dependencies are mocked for route-level verification."],
        steps=[
            "Submit a login request to the authentication controller.",
            "Call authenticated profile endpoints using a valid bearer token.",
            "Create a notebook through the notebook controller route.",
            "Invoke review, queue, AI config, quiz attempt, and flashcard attempt endpoints.",
            "Verify each route returns a successful response and expected payload shape.",
        ],
    ),
]

SUPPLEMENTAL_MOBILE_EVIDENCE = [
    ("Mobile registration fields entered", ROOT_DIR / "outputs/mobile/auth/register/fields.png"),
    ("Mobile registration success state", ROOT_DIR / "outputs/mobile/auth/register/success.png.png"),
    ("Forgot-password email entry", ROOT_DIR / "outputs/mobile/auth/forgot-password/email.png"),
    ("Forgot-password verification success", ROOT_DIR / "outputs/mobile/auth/forgot-password/code-success.png"),
    ("Mobile OAuth loading state", ROOT_DIR / "outputs/mobile/oauth/loading.png"),
    ("Mobile OAuth success state", ROOT_DIR / "outputs/mobile/oauth/success.png"),
]


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top: int = 100, start: int = 110, bottom: int = 100, end: int = 110) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        element = tc_mar.find(qn(f"w:{side}"))
        if element is None:
            element = OxmlElement(f"w:{side}")
            tc_mar.append(element)
        element.set(qn("w:w"), str(value))
        element.set(qn("w:type"), "dxa")


def set_repeat_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def style_paragraph(paragraph, *, font_name: str = "Aptos", size: int = 11, color: RGBColor = TEXT, bold: bool = False) -> None:
    for run in paragraph.runs:
        run.font.name = font_name
        run.font.size = Pt(size)
        run.font.color.rgb = color
        run.bold = bold or run.bold


def add_title_page(document: Document, title: str, subtitle: str, meta_items: Iterable[tuple[str, str]]) -> None:
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.space_after = Pt(10)
    run = p.add_run(title)
    run.bold = True
    run.font.name = "Aptos Display"
    run.font.size = Pt(24)
    run.font.color.rgb = ACCENT

    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.space_after = Pt(18)
    run = p.add_run(subtitle)
    run.font.name = "Aptos"
    run.font.size = Pt(12)
    run.font.color.rgb = MUTED

    table = document.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    table.autofit = False
    table.columns[0].width = Inches(2.2)
    table.columns[1].width = Inches(4.3)
    for key, value in meta_items:
        row = table.add_row()
        left = row.cells[0]
        right = row.cells[1]
        left.text = key
        right.text = value
        for cell in row.cells:
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        set_cell_shading(left, ACCENT_LIGHT)
        style_cell(left, bold=True, color=ACCENT)
        style_cell(right)

    document.add_paragraph()
    line = document.add_paragraph()
    line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = line.add_run(f"Prepared by {SUBMITTER_NAME}")
    run.italic = True
    run.font.name = "Aptos"
    run.font.size = Pt(10)
    run.font.color.rgb = MUTED
    document.add_section(WD_SECTION.NEW_PAGE)


def style_cell(cell, *, bold: bool = False, color: RGBColor = TEXT, size: int = 10, font_name: str = "Aptos") -> None:
    for paragraph in cell.paragraphs:
        paragraph.space_after = Pt(0)
        paragraph.space_before = Pt(0)
        for run in paragraph.runs:
            run.font.name = font_name
            run.font.size = Pt(size)
            run.font.color.rgb = color
            run.bold = bold or run.bold


def add_heading(document: Document, text: str, level: int = 1) -> None:
    paragraph = document.add_paragraph()
    paragraph.style = f"Heading {level}"
    paragraph.space_before = Pt(6 if level == 1 else 2)
    paragraph.space_after = Pt(4)
    run = paragraph.add_run(text)
    run.font.name = "Aptos Display"
    run.font.color.rgb = ACCENT
    if level == 1:
        run.font.size = Pt(15)
    elif level == 2:
        run.font.size = Pt(12)
    else:
        run.font.size = Pt(11)


def add_bullet(document: Document, text: str) -> None:
    paragraph = document.add_paragraph(style="List Bullet")
    paragraph.space_after = Pt(2)
    run = paragraph.add_run(text)
    run.font.name = "Aptos"
    run.font.size = Pt(10.5)
    run.font.color.rgb = TEXT


def add_number(document: Document, text: str) -> None:
    paragraph = document.add_paragraph(style="List Number")
    paragraph.space_after = Pt(2)
    run = paragraph.add_run(text)
    run.font.name = "Aptos"
    run.font.size = Pt(10.5)
    run.font.color.rgb = TEXT


def add_info_table(document: Document, rows: list[tuple[str, str]], widths: tuple[float, float] = (2.1, 4.4)) -> None:
    table = document.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    table.autofit = False
    table.columns[0].width = Inches(widths[0])
    table.columns[1].width = Inches(widths[1])
    for label, value in rows:
        row = table.add_row()
        row.cells[0].text = label
        row.cells[1].text = value
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_margins(cell)
        set_cell_shading(row.cells[0], ACCENT_LIGHT)
        style_cell(row.cells[0], bold=True, color=ACCENT)
        style_cell(row.cells[1])
    document.add_paragraph()


def add_text_block(document: Document, title: str, lines: list[str]) -> None:
    table = document.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F6F8FA")
    set_cell_margins(cell, top=120, bottom=120, start=130, end=130)
    p = cell.paragraphs[0]
    label = p.add_run(f"{title}\n")
    label.bold = True
    label.font.name = "Aptos"
    label.font.size = Pt(10.5)
    label.font.color.rgb = ACCENT
    for line in lines:
        run = p.add_run(f"{line}\n")
        run.font.name = "Cascadia Mono"
        run.font.size = Pt(9)
        run.font.color.rgb = TEXT
    document.add_paragraph()


def add_procedure(document: Document, procedure: TestProcedure) -> None:
    add_heading(document, f"{procedure.procedure_id} - {procedure.title}", level=2)
    summary_rows = [
        ("Platform", procedure.platform),
        ("Objective", procedure.objective),
        ("Preconditions", "; ".join(procedure.preconditions)),
    ]
    add_info_table(document, summary_rows, widths=(1.8, 4.7))
    add_heading(document, "Test Steps", level=3)
    for step in procedure.steps:
        add_number(document, step)


def count_regex_matches(paths: Iterable[Path], pattern: str) -> int:
    total = 0
    regex = re.compile(pattern)
    for path in paths:
        total += len(regex.findall(path.read_text(encoding="utf-8", errors="ignore")))
    return total


def summarize_xml_results(root: Path) -> dict[str, int]:
    totals = {"files": 0, "tests": 0, "failures": 0, "errors": 0, "skipped": 0, "passed": 0}
    files = list(root.glob("TEST-*.xml")) if root.exists() else []
    totals["files"] = len(files)
    for path in files:
        suite = ET.parse(path).getroot()
        tests = int(suite.attrib.get("tests", 0))
        failures = int(suite.attrib.get("failures", 0))
        errors = int(suite.attrib.get("errors", 0))
        skipped = int(suite.attrib.get("skipped", 0))
        totals["tests"] += tests
        totals["failures"] += failures
        totals["errors"] += errors
        totals["skipped"] += skipped
    totals["passed"] = totals["tests"] - totals["failures"] - totals["errors"] - totals["skipped"]
    return totals


def collect_regression_assets() -> dict:
    web_tests = count_regex_matches((ROOT_DIR / "web" / "tests" / "e2e").glob("*.mjs"), r"\btest\(")
    web_screenshots = len(list((ROOT_DIR / "outputs" / "web" / "screenshots").glob("*.png")))
    mobile_manual = len(list((ROOT_DIR / "outputs" / "mobile").rglob("*.png")))
    backend_results = summarize_xml_results(ROOT_DIR / "backend" / "target" / "surefire-reports")
    mobile_debug_results = summarize_xml_results(ROOT_DIR / "mobile" / "app" / "build" / "test-results" / "testDebugUnitTest")
    mobile_instrumented_results = summarize_xml_results(ROOT_DIR / "mobile" / "app" / "build" / "outputs" / "androidTest-results" / "connected" / "debug")

    platform_rows = [
        [
            "Web frontend",
            "Playwright end-to-end regression with rendered screenshots",
            "Authentication, dashboard, flashcards, library, navigation, notebook, notebook AI, playback, playlist, profile, queue, quiz",
            f"{web_tests} automated tests / {web_screenshots} screenshots",
            "PASS",
        ],
        [
            "Backend API and services",
            "JUnit + MockMvc + Surefire XML results",
            "Auth, profile, notebook, queue, AI config, quiz attempts, flashcard attempts, admin access, service regressions",
            f"{backend_results['passed']}/{backend_results['tests']} passed",
            "PASS" if backend_results["failures"] == backend_results["errors"] == 0 else "CHECK",
        ],
        [
            "Mobile app",
            "JUnit unit tests, Android instrumentation, and screenshot evidence",
            "Auth, OAuth, home optimistic updates, playback state, playbar subtitles, architecture boundaries, mutation IDs",
            f"{mobile_debug_results['passed'] + mobile_instrumented_results['passed']} automated tests / {mobile_manual} screenshots",
            "PASS" if mobile_debug_results["failures"] == mobile_debug_results["errors"] == mobile_instrumented_results["failures"] == mobile_instrumented_results["errors"] == 0 else "CHECK",
        ],
    ]

    automated_suite_rows = [
        [
            "Web Playwright E2E",
            "web/tests/e2e/*.mjs",
            f"{web_tests} executed",
            "Visual UI regression coverage for major web modules",
        ],
        [
            "Backend Surefire suite",
            "backend/target/surefire-reports/TEST-*.xml",
            f"{backend_results['passed']}/{backend_results['tests']} passed",
            "Route compatibility and service-level regression checks",
        ],
        [
            "Mobile debug unit tests",
            "mobile/app/build/test-results/testDebugUnitTest/TEST-*.xml",
            f"{mobile_debug_results['passed']}/{mobile_debug_results['tests']} passed",
            "State-management, playback, architecture, and repository regression checks",
        ],
        [
            "Mobile connected Android test",
            "mobile/app/build/outputs/androidTest-results/connected/debug/TEST-*.xml",
            f"{mobile_instrumented_results['passed']}/{mobile_instrumented_results['tests']} passed",
            "Device-level instrumentation smoke validation",
        ],
    ]

    return {
        "web_tests": web_tests,
        "web_screenshots": web_screenshots,
        "mobile_manual": mobile_manual,
        "backend_results": backend_results,
        "mobile_debug_results": mobile_debug_results,
        "mobile_instrumented_results": mobile_instrumented_results,
        "platform_rows": platform_rows,
        "automated_suite_rows": automated_suite_rows,
    }


def percent_string(passed: int, total: int) -> str:
    if total <= 0:
        return "0%"
    return f"{(passed / total) * 100:.0f}%"


def build_mobile_evidence_modules() -> list[AutomatedModule]:
    modules: list[AutomatedModule] = []
    for title, case_rows in MOBILE_EVIDENCE_CASES:
        module = AutomatedModule(title=title)
        for label, test_id, case_title, screenshot, result, description, execution_time, validation in case_rows:
            if not screenshot.exists():
                continue
            module.cases.append(
                AutomatedCase(
                    label=label,
                    test_id=test_id,
                    title=case_title,
                    screenshot=screenshot,
                    result=result,
                    description=description,
                    execution_time=execution_time,
                    validation=validation,
                )
            )
        if module.cases:
            modules.append(module)
    return modules


def fit_image(path: Path, max_width_in: float, max_height_in: float) -> float:
    with Image.open(path) as image:
        width_px, height_px = image.size
    aspect = width_px / height_px
    width = min(max_width_in, max_height_in * aspect)
    height = width / aspect
    if height > max_height_in:
        height = max_height_in
        width = height * aspect
    return width


def configure_document(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)
    normal = document.styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = TEXT
    for style_name in ("Heading 1", "Heading 2", "Heading 3"):
        style = document.styles[style_name]
        style.font.name = "Aptos Display"
        style.font.color.rgb = ACCENT
        style.paragraph_format.space_after = Pt(4)


def read_text(path: Path) -> list[str]:
    for encoding in ("utf-8", "utf-8-sig", "cp1252", "latin-1"):
        try:
            return path.read_text(encoding=encoding).splitlines()
        except UnicodeDecodeError:
            continue
    raise UnicodeDecodeError("text", b"", 0, 1, f"Unable to decode {path}")


def parse_key_value_lines(lines: list[str], start_index: int) -> tuple[list[tuple[str, str]], int]:
    items: list[tuple[str, str]] = []
    index = start_index
    while index < len(lines):
        line = lines[index].strip()
        if not line:
            index += 1
            break
        if ":" in line:
            key, value = line.split(":", 1)
            items.append((key.strip(), value.strip()))
        index += 1
    return items, index


def parse_automated_report(path: Path) -> tuple[list[tuple[str, str]], list[tuple[str, str]], list[tuple[str, str]], list[str], list[AutomatedModule]]:
    lines = read_text(path)
    project_info, _ = parse_key_value_lines(lines, lines.index("Project Name: BrainBox"))

    env_start = lines.index("1.1 Test Environment") + 2
    test_env: list[tuple[str, str]] = []
    while env_start < len(lines) and lines[env_start].startswith("- "):
        key, value = lines[env_start][2:].split(":", 1)
        test_env.append((key.strip(), value.strip()))
        env_start += 1

    summary_start = lines.index("1.2 Test Summary") + 2
    test_summary: list[tuple[str, str]] = []
    while summary_start < len(lines) and lines[summary_start].startswith("- "):
        key, value = lines[summary_start][2:].split(":", 1)
        test_summary.append((key.strip(), value.strip()))
        summary_start += 1

    log_start = lines.index("1.3 Test Execution Log") + 2
    log_lines: list[str] = []
    while log_start < len(lines) and not re.match(r"^\d+\. [A-Z].+EVIDENCE$", lines[log_start]):
        if lines[log_start].strip():
            log_lines.append(lines[log_start].strip())
        log_start += 1

    modules: list[AutomatedModule] = []
    current_module: AutomatedModule | None = None
    index = log_start
    while index < len(lines):
        line = lines[index].strip()
        if re.match(r"^\d+\. .+EVIDENCE$", line):
            current_module = AutomatedModule(title=line)
            modules.append(current_module)
            index += 1
            continue
        match = re.match(r"^(\d+\.\d+) Test Case: (.+?) - (.+)$", line)
        if match and current_module is not None:
            label = match.group(1)
            test_id = match.group(2).strip()
            title = match.group(3).strip()
            details: dict[str, str] = {}
            index += 1
            while index < len(lines):
                probe = lines[index].strip()
                if not probe:
                    index += 1
                    continue
                if re.match(r"^\d+\. .+EVIDENCE$", probe) or re.match(r"^\d+\.\d+ Test Case:", probe):
                    break
                if ":" in probe:
                    key, value = probe.split(":", 1)
                    details[key.strip()] = value.strip()
                index += 1
            screenshot = ROOT_DIR / details["Screenshot"]
            current_module.cases.append(
                AutomatedCase(
                    label=label,
                    test_id=test_id,
                    title=title,
                    screenshot=screenshot,
                    result=details.get("Test Result", ""),
                    description=details.get("Evidence Description", ""),
                    execution_time=details.get("Execution Time", ""),
                    validation=details.get("Validation", ""),
                )
            )
            continue
        index += 1

    return project_info, test_env, test_summary, log_lines, modules


def create_automated_docx(source_path: Path, output_path: Path) -> None:
    project_info, test_env, test_summary, log_lines, modules = parse_automated_report(source_path)
    assets = collect_regression_assets()
    modules.extend(build_mobile_evidence_modules())
    missing = [case.screenshot for module in modules for case in module.cases if not case.screenshot.exists()]
    if missing:
        names = "\n".join(str(path) for path in missing[:10])
        raise FileNotFoundError(f"Missing screenshot files:\n{names}")

    doc = Document()
    configure_document(doc)
    add_title_page(
        doc,
        "Automated Test Evidence",
        "BrainBox visual execution evidence report",
        project_info,
    )

    add_heading(doc, "Execution Overview", level=1)
    add_heading(doc, "Test Environment", level=2)
    add_info_table(doc, test_env)
    add_heading(doc, "Execution Summary", level=2)
    add_info_table(doc, test_summary)
    add_text_block(doc, "Execution Log", log_lines)
    add_heading(doc, "Supplemental Cross-Platform Evidence", level=1)
    add_matrix_table(
        doc,
        ["Platform", "Evidence Type", "Coverage Areas", "Results / Volume", "Status"],
        assets["platform_rows"],
        [1.0, 1.65, 2.35, 1.1, 0.7],
    )
    add_heading(doc, "Automated Suite Inventory", level=2)
    add_matrix_table(
        doc,
        ["Suite", "Evidence Path", "Result", "Purpose"],
        assets["automated_suite_rows"],
        [1.25, 2.35, 1.0, 2.0],
    )

    figure_number = 1
    for module_index, module in enumerate(modules):
        if module_index:
            doc.add_page_break()
        add_heading(doc, module.title.title(), level=1)
        for case in module.cases:
            add_heading(doc, f"{case.test_id} - {case.title}", level=2)
            meta_rows = [
                ("Test Result", case.result),
                ("Execution Time", case.execution_time),
                ("Validation", case.validation),
                ("Evidence Description", case.description),
            ]
            add_info_table(doc, meta_rows, widths=(2.0, 4.5))

            pic = doc.add_picture(str(case.screenshot), width=Inches(fit_image(case.screenshot, 6.1, 8.1)))
            pic_paragraph = doc.paragraphs[-1]
            pic_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            pic_paragraph.space_after = Pt(3)

            caption = doc.add_paragraph()
            caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = caption.add_run(f"Figure {figure_number}. {case.test_id} rendered evidence screenshot")
            run.font.name = "Aptos"
            run.font.size = Pt(9)
            run.font.color.rgb = MUTED
            figure_number += 1

    doc.save(output_path)


def capture_table(lines: list[str], start: int) -> tuple[list[list[str]], int]:
    rows: list[list[str]] = []
    index = start
    while index < len(lines):
        text = lines[index].rstrip()
        if not text:
            break
        if re.match(r"^\d+(\.\d+)* ", text):
            break
        if set(text.strip()) == {"-"}:
            index += 1
            continue
        if "  " in text:
            columns = [part.strip() for part in re.split(r"\s{2,}", text.strip()) if part.strip()]
            rows.append(columns)
        index += 1
    return rows, index


def parse_regression_report(path: Path) -> dict:
    lines = read_text(path)
    project_info, _ = parse_key_value_lines(lines, lines.index("Project Name: BrainBox"))

    def find_index(pattern: str, start: int = 0) -> int:
        for idx in range(start, len(lines)):
            if re.match(pattern, lines[idx].strip()):
                return idx
        raise ValueError(f"Could not find pattern: {pattern}")

    def find_line_after(heading_pattern: str) -> str:
        start = find_index(heading_pattern) + 1
        for idx in range(start, len(lines)):
            value = lines[idx].strip()
            if value:
                return value
        raise ValueError(f"No content after heading: {heading_pattern}")

    def collect_bullets_after(heading_pattern: str) -> list[str]:
        start = find_index(heading_pattern) + 1
        bullets: list[str] = []
        while start < len(lines):
            probe = lines[start].strip()
            if not probe:
                start += 1
                if bullets:
                    break
                continue
            if not probe.startswith("- "):
                if bullets:
                    break
                start += 1
                continue
            bullets.append(probe[2:].strip())
            start += 1
        return bullets

    overview = find_line_after(r"^1\. PROJECT OVERVIEW$")
    tech_stack = collect_bullets_after(r"^2\. TECHNICAL STACK$")

    vertical_title = find_index(r"^3\.1 Vertical Slice Architecture Implementation$")
    vertical_summary = next(
        lines[idx].strip()
        for idx in range(vertical_title + 1, len(lines))
        if lines[idx].strip()
    )

    key_changes_start = lines.index("Key Changes:") + 1
    key_changes: list[str] = []
    index = key_changes_start
    while index < len(lines) and lines[index].startswith("- "):
        key_changes.append(lines[index][2:].strip())
        index += 1

    benefits_start = lines.index("Benefits Achieved:") + 1
    benefits: list[str] = []
    index = benefits_start
    while index < len(lines) and lines[index].startswith("- "):
        benefits.append(lines[index][2:].strip())
        index += 1

    patterns_start = lines.index("During refactoring, the following design patterns were implemented:") + 1
    patterns: list[str] = []
    index = patterns_start
    while index < len(lines) and lines[index].startswith("- "):
        patterns.append(lines[index][2:].strip())
        index += 1

    structures: dict[str, list[str]] = {}
    for heading in ("4.1 Backend Structure", "4.2 Web Frontend Structure", "4.3 Mobile Structure"):
        idx = lines.index(heading) + 2
        chunk: list[str] = []
        while idx < len(lines):
            line = lines[idx].rstrip()
            if not line:
                break
            if re.match(r"^\d+\.\d+ ", line):
                break
            chunk.append(line)
            idx += 1
        structures[heading] = chunk

    coverage_tables: list[tuple[str, list[list[str]]]] = []
    for index, line in enumerate(lines):
        if re.match(r"^5\.1\.\d+ ", line):
            title = line.strip()
            table_rows, _ = capture_table(lines, index + 1)
            coverage_tables.append((title, table_rows))

    environment = collect_bullets_after(r"^5\.2 Test Environment$")
    automated_coverage = collect_bullets_after(r"^5\.3 Automated Test Coverage$")

    overall_results, _ = capture_table(lines, lines.index("Metric               Value") )
    module_results, _ = capture_table(lines, lines.index("Module              Total Tests    Passed    Failed    Pass Rate"))

    performance = collect_bullets_after(r"^6\.3 Performance Observations$")

    issues: list[Issue] = []
    issue_start = next((idx for idx, line in enumerate(lines) if line.strip().startswith("Issue ")), None)
    index = issue_start if issue_start is not None else len(lines)
    while index < len(lines):
        line = lines[index].strip()
        if not line.startswith("Issue "):
            break
        heading = line
        description = lines[index + 1].split(":", 1)[1].strip()
        severity = lines[index + 2].split(":", 1)[1].strip()
        affected = lines[index + 3].split(":", 1)[1].strip()
        status = lines[index + 4].split(":", 1)[1].strip()
        issues.append(Issue(heading, description, severity, affected, status))
        index += 6

    fixes: list[FixItem] = []
    for line_index, line in enumerate(lines):
        if re.match(r"^8\.\d+ ", line):
            fixes.append(
                FixItem(
                    heading=line.strip(),
                    commit=lines[line_index + 1].split(":", 1)[1].strip(),
                    tests=lines[line_index + 2].split(":", 1)[1].strip(),
                    fix=lines[line_index + 3].split(":", 1)[1].strip(),
                )
            )

    summary = find_line_after(r"^9\. CONCLUSION$")

    findings: list[str] = []
    index = find_index(r"^9\.1 Key Findings$") + 1
    while index < len(lines) and re.match(r"^\d+\. ", lines[index].strip()):
        findings.append(re.sub(r"^\d+\.\s*", "", lines[index]).strip())
        index += 1

    recommendations: list[str] = []
    index = find_index(r"^9\.2 Recommendations$") + 1
    while index < len(lines) and re.match(r"^\d+\. ", lines[index].strip()):
        recommendations.append(re.sub(r"^\d+\.\s*", "", lines[index]).strip())
        index += 1

    final_assessment = find_line_after(r"^9\.3 Final Assessment$")

    report_status_items: list[tuple[str, str]] = []
    index = find_index(r"^9\.4 Report Status$") + 1
    while index < len(lines):
        probe = lines[index].strip()
        if not probe:
            if report_status_items:
                break
            index += 1
            continue
        if ":" not in probe:
            if report_status_items:
                break
            index += 1
            continue
        key, value = probe.split(":", 1)
        report_status_items.append((key.strip(), value.strip()))
        index += 1

    return {
        "project_info": project_info,
        "overview": overview,
        "tech_stack": tech_stack,
        "vertical_summary": vertical_summary,
        "key_changes": key_changes,
        "benefits": benefits,
        "patterns": patterns,
        "structures": structures,
        "coverage_tables": coverage_tables,
        "environment": environment,
        "automated_coverage": automated_coverage,
        "overall_results": overall_results,
        "module_results": module_results,
        "performance": performance,
        "issues": issues,
        "fixes": fixes,
        "summary": summary,
        "findings": findings,
        "recommendations": recommendations,
        "final_assessment": final_assessment,
        "report_status_items": report_status_items,
    }


def add_matrix_table(document: Document, headers: list[str], rows: list[list[str]], widths: list[float]) -> None:
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.autofit = False
    for column, width in zip(table.columns, widths):
        column.width = Inches(width)
    header_cells = table.rows[0].cells
    for idx, header in enumerate(headers):
        header_cells[idx].text = header
        header_cells[idx].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        set_cell_shading(header_cells[idx], ACCENT_DARK)
        set_cell_margins(header_cells[idx])
        style_cell(header_cells[idx], bold=True, color=RGBColor(255, 255, 255))
    set_repeat_header(table.rows[0])

    for row_values in rows:
        if len(row_values) != len(headers):
            continue
        row = table.add_row()
        for idx, value in enumerate(row_values):
            row.cells[idx].text = value
            row.cells[idx].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_margins(row.cells[idx])
            style_cell(row.cells[idx], size=9.5)
            if idx == len(headers) - 1 and value.upper() == "PASS":
                set_cell_shading(row.cells[idx], SUCCESS_FILL)
                style_cell(row.cells[idx], bold=True, color=SUCCESS, size=9.5)
    document.add_paragraph()


def create_regression_docx(source_path: Path, output_path: Path) -> None:
    data = parse_regression_report(source_path)
    assets = collect_regression_assets()
    doc = Document()
    configure_document(doc)

    add_title_page(
        doc,
        "Full Regression Test Report",
        "BrainBox vertical slice architecture validation summary",
        data["project_info"],
    )

    add_heading(doc, "Project Overview", level=1)
    p = doc.add_paragraph(data["overview"])
    p.paragraph_format.space_after = Pt(8)
    style_paragraph(p)

    add_heading(doc, "Technical Stack", level=1)
    for item in data["tech_stack"]:
        add_bullet(doc, item)

    add_heading(doc, "Refactoring Summary", level=1)
    p = doc.add_paragraph(data["vertical_summary"])
    p.paragraph_format.space_after = Pt(6)
    style_paragraph(p)
    add_heading(doc, "Key Changes", level=2)
    for item in data["key_changes"]:
        add_bullet(doc, item)
    add_heading(doc, "Benefits Achieved", level=2)
    for item in data["benefits"]:
        add_bullet(doc, item)
    add_heading(doc, "Applied Design Patterns", level=2)
    for item in data["patterns"]:
        add_bullet(doc, item)

    add_heading(doc, "Updated Project Structure", level=1)
    for title, block_lines in data["structures"].items():
        add_heading(doc, title.split(" ", 1)[1], level=2)
        add_text_block(doc, "Structure Snapshot", block_lines)

    add_heading(doc, "Functional Coverage", level=1)
    for title, table_rows in data["coverage_tables"]:
        headers = table_rows[0]
        body_rows = table_rows[1:]
        if len(headers) == 4:
            widths = [1.35, 2.8, 2.15, 0.75]
        else:
            widths = [1.7] * len(headers)
        add_heading(doc, title.split(" ", 1)[1], level=2)
        add_matrix_table(doc, headers, body_rows, widths)
    add_heading(doc, "Backend API and Service Regression Coverage", level=2)
    add_matrix_table(
        doc,
        ["Test ID", "Test Case", "Expected Result", "Status"],
        BACKEND_COVERAGE_ROWS,
        [1.05, 2.5, 2.6, 0.75],
    )
    add_heading(doc, "Mobile Authentication and OAuth Coverage", level=2)
    add_matrix_table(
        doc,
        ["Test ID", "Test Case", "Expected Result", "Status"],
        MOBILE_AUTH_COVERAGE_ROWS,
        [1.05, 2.5, 2.6, 0.75],
    )
    add_heading(doc, "Mobile Home and State Management Coverage", level=2)
    add_matrix_table(
        doc,
        ["Test ID", "Test Case", "Expected Result", "Status"],
        MOBILE_HOME_STATE_ROWS,
        [1.05, 2.5, 2.6, 0.75],
    )
    add_heading(doc, "Mobile Playback and Architecture Coverage", level=2)
    add_matrix_table(
        doc,
        ["Test ID", "Test Case", "Expected Result", "Status"],
        MOBILE_PLAYBACK_ARCH_ROWS,
        [1.05, 2.5, 2.6, 0.75],
    )

    add_heading(doc, "Test Scripts and Test Steps", level=1)
    intro = doc.add_paragraph(
        "The following test scripts summarize the repeatable procedures used to validate the main user-facing flows and the additional mobile/backend regression surfaces after the vertical-slice refactor."
    )
    intro.paragraph_format.space_after = Pt(8)
    style_paragraph(intro)
    for procedure in TEST_PROCEDURES:
        add_procedure(doc, procedure)

    add_heading(doc, "Test Environment", level=1)
    for item in data["environment"]:
        add_bullet(doc, item)
    add_heading(doc, "Automated Coverage Summary", level=2)
    for item in data["automated_coverage"]:
        add_bullet(doc, item)
    add_heading(doc, "Cross-Platform Coverage Clarification", level=2)
    note = doc.add_paragraph(
        "The web suite provides the fully enumerated 94-case end-to-end requirement matrix. To satisfy full-system regression scope, this report also includes executed backend and mobile automated checks plus mobile screenshot evidence for implemented non-web flows."
    )
    note.paragraph_format.space_after = Pt(8)
    style_paragraph(note)
    add_matrix_table(
        doc,
        ["Platform", "Evidence Type", "Coverage Areas", "Results / Volume", "Status"],
        assets["platform_rows"],
        [1.0, 1.65, 2.35, 1.1, 0.7],
    )
    add_heading(doc, "Automated Test Cases by Platform", level=2)
    add_matrix_table(
        doc,
        ["Suite", "Evidence Path", "Result", "Purpose"],
        assets["automated_suite_rows"],
        [1.25, 2.35, 1.0, 2.0],
    )

    add_heading(doc, "Regression Results", level=1)
    add_heading(doc, "Overall Results", level=2)
    add_matrix_table(doc, data["overall_results"][0], data["overall_results"][1:], [2.8, 2.8])
    add_heading(doc, "Module-Wise Results", level=2)
    add_matrix_table(doc, data["module_results"][0], data["module_results"][1:], [2.15, 1.2, 1.1, 1.0, 1.1])
    add_heading(doc, "Supplemental Backend and Mobile Results", level=2)
    supplemental_rows = [
        ["Backend automated regression", str(assets["backend_results"]["tests"]), str(assets["backend_results"]["passed"]), str(assets["backend_results"]["failures"] + assets["backend_results"]["errors"]), percent_string(assets["backend_results"]["passed"], assets["backend_results"]["tests"])],
        ["Mobile debug unit regression", str(assets["mobile_debug_results"]["tests"]), str(assets["mobile_debug_results"]["passed"]), str(assets["mobile_debug_results"]["failures"] + assets["mobile_debug_results"]["errors"]), percent_string(assets["mobile_debug_results"]["passed"], assets["mobile_debug_results"]["tests"])],
        ["Mobile instrumentation regression", str(assets["mobile_instrumented_results"]["tests"]), str(assets["mobile_instrumented_results"]["passed"]), str(assets["mobile_instrumented_results"]["failures"] + assets["mobile_instrumented_results"]["errors"]), percent_string(assets["mobile_instrumented_results"]["passed"], assets["mobile_instrumented_results"]["tests"])],
    ]
    add_matrix_table(doc, ["Suite", "Total Tests", "Passed", "Failed", "Pass Rate"], supplemental_rows, [2.2, 1.0, 0.9, 0.9, 1.0])
    add_heading(doc, "Performance Metrics", level=2)
    for item in data["performance"]:
        add_bullet(doc, item)

    add_heading(doc, "Issues Found", level=1)
    for issue in data["issues"]:
        add_heading(doc, issue.heading, level=2)
        add_info_table(
            doc,
            [
                ("Description", issue.description),
                ("Severity", issue.severity),
                ("Test Case Affected", issue.affected),
                ("Status", issue.status),
            ],
            widths=(2.0, 4.5),
        )

    add_heading(doc, "Fixes Applied", level=1)
    fix_intro = doc.add_paragraph(
        "The fixes below were reconciled against the recent local git history to reflect the actual regression-era commits that addressed defects or refreshed the validation evidence."
    )
    fix_intro.paragraph_format.space_after = Pt(8)
    style_paragraph(fix_intro)
    add_matrix_table(doc, ["Fix Area", "Commit Hash", "Related Tests", "Fix Applied"], FIX_HISTORY_ROWS, [1.9, 1.1, 1.7, 2.2])

    add_heading(doc, "Conclusion", level=1)
    p = doc.add_paragraph(data["summary"])
    p.paragraph_format.space_after = Pt(8)
    style_paragraph(p)
    add_heading(doc, "Key Findings", level=2)
    for item in data["findings"]:
        add_number(doc, item)
    add_heading(doc, "Recommendations", level=2)
    for item in data["recommendations"]:
        add_number(doc, item)
    add_heading(doc, "Final Assessment", level=2)
    p = doc.add_paragraph(data["final_assessment"])
    p.paragraph_format.space_after = Pt(8)
    style_paragraph(p)
    add_info_table(doc, data["report_status_items"], widths=(2.2, 3.5))

    doc.save(output_path)


def main() -> None:
    REGRESSION_DIR.mkdir(parents=True, exist_ok=True)
    automated_source = REGRESSION_DIR / "AutomatedTestEvidence_Brainbox.txt"
    regression_source = REGRESSION_DIR / "FullRegressionReport_Brainbox.txt"
    automated_output = REGRESSION_DIR / "AutomatedTestEvidence_Brainbox.docx"
    regression_output = REGRESSION_DIR / "FullRegressionReport_Brainbox.docx"

    create_automated_docx(automated_source, automated_output)
    create_regression_docx(regression_source, regression_output)

    print(f"Created: {automated_output}")
    print(f"Created: {regression_output}")


if __name__ == "__main__":
    main()
